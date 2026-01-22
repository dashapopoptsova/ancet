#!/usr/bin/env python3
"""
Minimal DOCX filler for "no-template" forms.

Input JSON formats supported:
1) {"field label 1": "value1", "field label 2": "value2", ...}
2) {"fields": [{"anchor": "...", "value": "..."}, ...]}

Heuristics:
- If anchor found in a table cell -> write to right cell (same row, next col).
- If anchor found in paragraph:
  - replace runs of underscores "_____" with value
  - else if ":" present -> set text after ":" to value
  - else if next paragraph is empty / underscores -> fill next paragraph
  - else append value at end
"""
from __future__ import annotations

import argparse
import json
import re
from dataclasses import dataclass
from difflib import SequenceMatcher
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph


UNDERSCORE_RE = re.compile(r"_{3,}")  # 3+ underscores treated as a blank
WS_RE = re.compile(r"\s+")


def norm(s: str) -> str:
    s = s.strip().lower()
    s = WS_RE.sub(" ", s)
    return s


def sim(a: str, b: str) -> float:
    """Rough similarity 0..1 (no extra deps)."""
    if not a or not b:
        return 0.0
    return SequenceMatcher(None, norm(a), norm(b)).ratio()


@dataclass
class Match:
    kind: str  # "paragraph" | "cell"
    score: float
    paragraph: Optional[Paragraph] = None
    table: Optional[Table] = None
    cell: Optional[_Cell] = None
    cell_pos: Optional[Tuple[int, int]] = None  # (row_idx, col_idx)


def load_fields(json_path: Path) -> List[Tuple[str, str]]:
    data = json.loads(json_path.read_text(encoding="utf-8"))

    # 1) { "fields": [ { "anchor": "...", "value": "..." }, ... ] }
    if isinstance(data, dict) and isinstance(data.get("fields"), list):
        out: List[Tuple[str, str]] = []
        for item in data["fields"]:
            if not isinstance(item, dict):
                continue
            anchor = str(item.get("anchor", "")).strip()
            value = item.get("value", "")
            value_str = "" if value is None else str(value).strip()
            if anchor:
                out.append((anchor, value_str))
        return out

    # 2) { "label": "value", ... }
    if isinstance(data, dict):
        return [(str(k).strip(), "" if v is None else str(v).strip()) for k, v in data.items()]

    # 3) YOUR CASE: [ {"Some label": "value"}, {"Another label": "value2"} ]
    if isinstance(data, list):
        out: List[Tuple[str, str]] = []
        for item in data:
            if not isinstance(item, dict) or not item:
                continue

            # if it's {"anchor": "...", "value": "..."} form
            if "anchor" in item and "value" in item:
                anchor = str(item["anchor"]).strip()
                value = item["value"]
                out.append((anchor, "" if value is None else str(value).strip()))
                continue

            # singleton dict form: {"<label>": "<value>"}
            if len(item) == 1:
                (k, v), = item.items()
                anchor = str(k).strip()
                value_str = "" if v is None else str(v).strip()
                out.append((anchor, value_str))
                continue

            # fallback: ignore unknown dict shapes
        return out

    raise ValueError(
        "Unsupported JSON format. Expected dict, {fields:[{anchor,value}]}, "
        "or list of singleton dicts [{label:value}, ...]."
    )

def iter_paragraphs(doc: Document) -> Iterable[Paragraph]:
    for p in doc.paragraphs:
        yield p


def iter_cells(doc: Document) -> Iterable[Tuple[Table, _Cell, int, int]]:
    for t_i, table in enumerate(doc.tables):
        for r_i, row in enumerate(table.rows):
            for c_i, cell in enumerate(row.cells):
                yield table, cell, r_i, c_i


def find_best_anchor(doc: Document, anchor: str, min_score: float = 0.62) -> Optional[Match]:
    anchor_n = norm(anchor)

    best: Optional[Match] = None

    # paragraphs
    for p in iter_paragraphs(doc):
        text = p.text or ""
        if not text.strip():
            continue
        score = sim(anchor_n, text)
        # boost if anchor is a substring
        if anchor_n and anchor_n in norm(text):
            score = max(score, 0.95)
        if best is None or score > best.score:
            best = Match(kind="paragraph", score=score, paragraph=p)

    # table cells
    for table, cell, r_i, c_i in iter_cells(doc):
        text = cell.text or ""
        if not text.strip():
            continue
        score = sim(anchor_n, text)
        if anchor_n and anchor_n in norm(text):
            score = max(score, 0.95)
        if best is None or score > best.score:
            best = Match(kind="cell", score=score, table=table, cell=cell, cell_pos=(r_i, c_i))

    if best is None or best.score < min_score:
        return None
    return best


def replace_underscores(text: str, value: str) -> str:
    if UNDERSCORE_RE.search(text):
        return UNDERSCORE_RE.sub(value if value else "—", text, count=1)
    return text


def fill_paragraph(doc: Document, p: Paragraph, value: str) -> bool:
    """Try to fill in-place; returns True if changed."""
    original = p.text or ""
    if not original.strip():
        return False

    # 1) underscores
    new_text = replace_underscores(original, value)
    if new_text != original:
        p.text = new_text
        return True

    # 2) colon pattern "Label: ____"
    if ":" in original:
        left, right = original.split(":", 1)
        # if right already has something meaningful, don't overwrite blindly
        right_clean = right.strip()
        if not right_clean or UNDERSCORE_RE.search(right) or right_clean in {"—", "-", "–"}:
            insert_val = value if value else "—"
            p.text = f"{left.strip()}: {insert_val}"
            return True

    # 3) next paragraph if empty/underscores
    # Find p index in doc.paragraphs
    try:
        idx = doc.paragraphs.index(p)
    except ValueError:
        idx = -1

    if idx >= 0 and idx + 1 < len(doc.paragraphs):
        p_next = doc.paragraphs[idx + 1]
        nxt = (p_next.text or "").strip()
        if (not nxt) or UNDERSCORE_RE.search(p_next.text or ""):
            p_next.text = replace_underscores(p_next.text or "", value if value else "—")
            if not (p_next.text or "").strip():
                p_next.text = value if value else "—"
            return True

    # 4) fallback: append at end (minimal)
    p.text = f"{original.rstrip()} {value if value else '—'}"
    return True


def fill_table_cell_right(match: Match, value: str) -> bool:
    assert match.table is not None and match.cell_pos is not None
    r_i, c_i = match.cell_pos
    row = match.table.rows[r_i]
    if c_i + 1 < len(row.cells):
        target = row.cells[c_i + 1]
        txt = target.text or ""
        # if underscores -> replace; else overwrite if empty
        if UNDERSCORE_RE.search(txt):
            target.text = replace_underscores(txt, value if value else "—")
        elif not txt.strip() or txt.strip() in {"—", "-", "–"}:
            target.text = value if value else "—"
        else:
            # already filled -> append (safe minimal)
            target.text = f"{txt.rstrip()} {value if value else '—'}"
        return True

    # if no right cell exists, write into the same cell (fallback)
    cell_txt = match.cell.text or ""
    match.cell.text = replace_underscores(cell_txt, value if value else "—")
    if match.cell.text == cell_txt:
        match.cell.text = f"{cell_txt.rstrip()} {value if value else '—'}"
    return True


def fill_docx(input_docx: Path, input_json: Path, output_docx: Path) -> None:
    doc = Document(str(input_docx))
    fields = load_fields(input_json)

    not_found: List[str] = []

    for anchor, value in fields:
        if not anchor:
            continue
        m = find_best_anchor(doc, anchor)
        if m is None:
            not_found.append(anchor)
            continue

        if m.kind == "cell":
            fill_table_cell_right(m, value)
        else:
            assert m.paragraph is not None
            fill_paragraph(doc, m.paragraph, value)

    doc.save(str(output_docx))

    # Minimal log to stdout (optional)
    if not_found:
        print("Anchors not found (check labels or lower similarity threshold):")
        for a in not_found[:50]:
            print(" -", a)


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("docx", type=Path, nargs="?", default=Path("АнкетаБенефициара.docx"))
    ap.add_argument("json", type=Path, nargs="?", default=Path("filled_fields.json"))
    ap.add_argument("out", type=Path, nargs="?", default=Path("result.docx"))
    args = ap.parse_args()

    fill_docx(args.docx, args.json, args.out)
    return 0

if __name__ == "__main__":
    raise SystemExit(main())

